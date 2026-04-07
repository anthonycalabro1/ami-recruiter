"""
Resume Parser - Extracts text from PDF/Word/TXT files and uses Claude API
to parse into structured candidate profiles.
"""

import os
import json
import anthropic
from pathlib import Path
from config_loader import CONFIG
from api_utils import retry_api_call


def extract_text_from_file(filepath):
    """Extract text content from PDF, DOCX, or TXT files."""
    ext = Path(filepath).suffix.lower()

    if ext == '.pdf':
        return _extract_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return _extract_docx(filepath)
    elif ext == '.txt':
        return _extract_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(filepath):
    """Extract text from PDF using pdfplumber."""
    import pdfplumber
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def _extract_docx(filepath):
    """Extract text from Word document.

    Uses raw XML parsing to capture all text, including content inside
    nested tables and complex layouts that python-docx's high-level API misses.
    Falls back to the high-level API result if XML extraction yields less text.
    """
    import docx

    # High-level extraction (paragraphs + top-level tables)
    hl_text = ""
    try:
        doc = docx.Document(filepath)
        hl_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                hl_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    hl_parts.append(" | ".join(row_texts))
        hl_text = "\n".join(hl_parts)
    except Exception:
        # python-docx can fail on non-standard DOCX files (missing
        # relationships, etc.).  Fall through to raw XML extraction.
        pass

    # Raw XML extraction — catches nested tables, content controls, text boxes,
    # and works even when python-docx rejects the file.
    import zipfile
    import xml.etree.ElementTree as ET
    # Transitional and Strict OOXML use different namespaces
    ns_transitional = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    ns_strict = "{http://purl.oclc.org/ooxml/wordprocessingml/main}"
    xml_parts = []
    current_para = []
    try:
        with zipfile.ZipFile(filepath) as zf:
            for part_name in zf.namelist():
                if part_name in ("word/document.xml", "word/header1.xml",
                                 "word/header2.xml", "word/footer1.xml",
                                 "word/footer2.xml"):
                    tree = ET.parse(zf.open(part_name))
                    for elem in tree.iter():
                        tag = elem.tag
                        if tag.endswith("}p") and (tag.startswith(ns_transitional) or tag.startswith(ns_strict)):
                            # Flush previous paragraph
                            para_text = "".join(current_para).strip()
                            if para_text:
                                xml_parts.append(para_text)
                            current_para = []
                        elif tag.endswith("}t") and (tag.startswith(ns_transitional) or tag.startswith(ns_strict)) and elem.text:
                            current_para.append(elem.text)
                        elif tag.endswith("}tab") and (tag.startswith(ns_transitional) or tag.startswith(ns_strict)):
                            current_para.append("\t")
                        elif tag.endswith("}br") and (tag.startswith(ns_transitional) or tag.startswith(ns_strict)):
                            current_para.append("\n")
            # Flush last paragraph
            para_text = "".join(current_para).strip()
            if para_text:
                xml_parts.append(para_text)
    except zipfile.BadZipFile:
        # Not a valid ZIP/DOCX at all — may be a legacy .doc binary format
        if hl_text:
            return hl_text
        raise ValueError(
            f"Cannot read {Path(filepath).name}: file is not a valid DOCX "
            f"(ZIP) archive. If this is a legacy .doc file, please re-save "
            f"it as .docx in Word first."
        )

    xml_text = "\n".join(xml_parts)

    if not xml_text and not hl_text:
        raise ValueError(
            f"Cannot extract text from {Path(filepath).name}: the file may be "
            f"a legacy .doc binary format. Please re-save it as .docx in Word."
        )

    # Return whichever method captured more content
    return xml_text if len(xml_text) >= len(hl_text) else hl_text


def _extract_txt(filepath):
    """Extract text from plain text file."""
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        return f.read().strip()


PARSE_PROMPT = """You are an expert AMI (Advanced Metering Infrastructure) recruiting analyst. Your job is to parse a resume and extract structured information relevant to AMI consulting roles.

Analyze the following resume and extract the information below. Be thorough and precise. If information is not found in the resume, use null.

IMPORTANT RULES:
- For AMI experience years: ONLY count time spent on metering-focused programs. The following terms are ALL synonymous with AMI and SHOULD be counted: "Advanced Metering Infrastructure", "AMI", "smart meter/smart metering", "next generation metering". General utility, smart grid, or grid modernization work without metering specificity does NOT count — the work must be focused on metering to qualify. Roles that explicitly reference AMI programs even without using the word "AMI" in the title DO count — e.g., deployment execution, hypercare, release governance, or issue triage on a program that is otherwise documented as AMI on the resume should be counted. When in doubt about a borderline role, include it if the surrounding context (same program, same client) is clearly AMI.
- For functional areas: Map the candidate's experience to these five areas: Strategy & Business Case, Business Integration, System Integration, Field Deployment Management, AMI Operations.
- A candidate may have experience in multiple functional areas.
- Look for specific AMI signals: vendor names (Itron, Landis+Gyr, Sensus/Xylem, Siemens/Trilliant), AMI-specific terminology (head-end, MDMS, VEE, OTA, mesh network, remote connect/disconnect, meter-to-cash, etc.), named AMI programs, meter deployment references, AND metering-synonymous terms (smart meter, smart metering, next generation metering).

Return your analysis as a JSON object with exactly this structure:

{
    "name": "Full Name",
    "email": "email@example.com or null",
    "phone": "phone number or null",
    "linkedin_url": "LinkedIn URL or null",
    "total_years_experience": 15,
    "total_ami_years": 7,
    "ami_years_explanation": "Brief explanation of how you calculated AMI years, listing specific programs/roles counted",
    "consulting_vs_utility": "consulting" | "utility" | "mixed" | "vendor",
    "consulting_detail": "Brief description of consulting vs utility vs vendor experience breakdown",
    "current_role": "Current or most recent title and company",
    "education": "Highest degree and institution",
    "certifications": ["PMP", "etc"],
    "ami_programs": [
        {
            "client_or_employer": "Utility name or consulting firm",
            "program_description": "Brief description",
            "role": "Their role on this program",
            "duration_years": 2.5,
            "meter_scale": "Number of meters if mentioned",
            "ami_vendors_used": ["Itron", "Landis+Gyr"],
            "functional_areas_touched": ["Business Integration", "System Integration"]
        }
    ],
    "ami_vendor_platforms": ["Itron IEE", "MV-RS", "Gridstream", "etc"],
    "ami_technologies_mentioned": ["head-end", "MDMS", "VEE", "mesh network", "cellular", "etc"],
    "ami_specific_functionality": ["remote connect/disconnect", "OTA firmware", "on-demand reads", "etc"],
    "programming_languages": ["Python", "SQL", "etc"],
    "functional_area_assessment": {
        "strategy_business_case": {
            "has_experience": true/false,
            "years_in_area": 3,
            "ami_specific": true/false,
            "key_activities": ["business case development", "regulatory filing support", "etc"],
            "summary": "2-3 sentence assessment of their experience in this area"
        },
        "business_integration": {
            "has_experience": true/false,
            "years_in_area": 4,
            "ami_specific": true/false,
            "subcategories": {
                "business_process_design": {"has_experience": true/false, "summary": "brief"},
                "project_management": {"has_experience": true/false, "summary": "brief"},
                "change_management": {"has_experience": true/false, "summary": "brief"}
            },
            "key_activities": ["BPD workshops", "process flows", "etc"],
            "summary": "2-3 sentence assessment"
        },
        "system_integration": {
            "has_experience": true/false,
            "years_in_area": 2,
            "ami_specific": true/false,
            "subcategories": {
                "solution_architecture": {"has_experience": true/false, "summary": "brief"},
                "technical_design": {"has_experience": true/false, "summary": "brief"},
                "applications_development": {"has_experience": true/false, "summary": "brief"},
                "testing": {"has_experience": true/false, "summary": "brief"},
                "cutover": {"has_experience": true/false, "summary": "brief"}
            },
            "key_activities": ["integration development", "MDMS configuration", "etc"],
            "summary": "2-3 sentence assessment"
        },
        "field_deployment_management": {
            "has_experience": true/false,
            "years_in_area": 1,
            "ami_specific": true/false,
            "key_activities": ["deployment strategy", "MIC management", "etc"],
            "mic_vendors_named": ["Aclara", "UPA", "etc"],
            "summary": "2-3 sentence assessment"
        },
        "ami_operations": {
            "has_experience": true/false,
            "years_in_area": 2,
            "ami_specific": true/false,
            "has_production_experience": true/false,
            "key_activities": ["VEE management", "meter provisioning", "etc"],
            "summary": "2-3 sentence assessment"
        }
    },
    "overall_assessment": "3-4 sentence overall assessment of this candidate's AMI experience profile, strengths, and potential gaps",
    "red_flags": ["Any concerns about the resume, such as vague AMI references, grid modernization without AMI specificity, etc."]
}

RESUME TEXT:
"""


def parse_resume(resume_text):
    """Send resume text to Claude API for structured parsing."""
    client = anthropic.Anthropic(api_key=CONFIG['anthropic_api_key'])

    def _call():
        return client.messages.create(
            model=CONFIG['model'],
            max_tokens=4096,
            messages=[{"role": "user", "content": PARSE_PROMPT + resume_text}]
        )

    message = retry_api_call(_call)

    response_text = message.content[0].text

    # Extract JSON from response (handle potential markdown code blocks)
    if "```json" in response_text:
        response_text = response_text.split("```json")[1].split("```")[0]
    elif "```" in response_text:
        response_text = response_text.split("```")[1].split("```")[0]

    try:
        parsed = json.loads(response_text.strip())
    except json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse Claude's response as JSON: {e}\nResponse: {response_text[:500]}")

    return parsed


def determine_role_routing(ami_years):
    """Determine role routing based on AMI years of experience."""
    if ami_years is None or ami_years < 3:
        return "eliminated"
    elif ami_years <= 5:
        return "senior_only"
    elif ami_years <= 7:
        return "senior_plus_manager_flag"
    else:
        return "manager_only"


def get_matching_functional_areas(parsed_profile):
    """Identify which functional areas the candidate should be scored against."""
    matching = []
    fa = parsed_profile.get('functional_area_assessment', {})

    area_map = {
        'strategy_business_case': 'Strategy & Business Case',
        'business_integration': 'Business Integration',
        'system_integration': 'System Integration',
        'field_deployment_management': 'Field Deployment Management',
        'ami_operations': 'AMI Operations'
    }

    for key, display_name in area_map.items():
        area_data = fa.get(key, {})
        if area_data.get('has_experience') and area_data.get('ami_specific'):
            matching.append(display_name)

    return matching
